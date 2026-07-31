from pathlib import Path
import fitz
import csv
from docx import Document

from langchain_text_splitters import RecursiveCharacterTextSplitter
from rag.embedding import embed
from rag.vectordb import get_collection, delete_by_source
from rag.bm25_store import save_bm25
from rag.hash_store import get_changed_files, save_hashes


# Config
DATA_PATH = "data/raw"
CHUNK_SIZE = 500
OVERLAP = 100

# PDF Reader
def read_pdf(file_path):
    doc = fitz.open(file_path)
    text = ""

    for page in doc:
        page_text = page.get_text()
        if page_text:
            text += page_text + "\n"

    return text

# DOCX Reader
def read_docx(file_path):
    """Extract text from Word document preserving paragraph structure."""
    doc  = Document(file_path)
    text = ""
    for para in doc.paragraphs:
        if para.text.strip():
            text += para.text + "\n"
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                text += row_text + "\n"
    return text

# CSV Reader
def read_csv(file_path):
    text= ""
    with open(file_path, "r", encoding = "utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            row_text = ", ".join(f"{k}: {v}" for k, v in row.items() if v.strip())
            text += row_text + "\n"
    return text


# Load Files
def load_files(path):
    texts = []

    for file in Path(path).glob("*"):
        print(f"Loading:{file.name}")
        suffix = file.suffix.lower()

        if suffix == ".txt":
            with open(file, "r", encoding="utf-8") as f:
                texts.append((file.name,f.read(), str(file)))

        elif suffix == ".pdf":
            text = read_pdf(file)
            texts.append((file.name,text, str(file)))
        
        elif suffix == ".csv":
            text = read_csv(file)
            texts.append((file.name, text, str(file)))

        elif suffix == ".docx":
            text = read_docx(file)
            texts.append((file.name, text, str(file)))

    return texts


# Chunking(Langchain)
splitter = RecursiveCharacterTextSplitter(
    chunk_size = CHUNK_SIZE,
    chunk_overlap = OVERLAP
)

# PDF (split by paragraph, first then by size)
def chunk_pdf(text):
    # Split on double newline (paragraph) first
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunks = []
    for para in paragraphs:
        if len(para) > CHUNK_SIZE:
            # Split large paragraphs further
            chunks.extend(splitter.split_text(para))
        else:
            chunks.append(para)
    return chunks
    
# CSV (each row is already a natural chunk , dont split rows)
def chunk_csv(text):
    # Each line in one record - keep them together
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    chunks = []
    current = ""
    for line in lines:
        if len(current) + len(line) < CHUNK_SIZE:
            current += line + "\n"
        else:
            if current:
                chunks.append(current.strip())
            current = line + "\n"
    if current:
        chunks.append(current.strip())
    return chunks

# TXT (use standard recursive splitter)
def chunk_text(text):
    return splitter.split_text(text)


# Ingest Pipeline
def ingest():
    collection = get_collection()
    all_files = load_files(DATA_PATH)

    if not all_files:
        print("⚠️  No files found in data/raw")
        return

    # Hash check
    filepaths = [fp for _, _, fp in all_files]
    changed_paths, updated_hashes = get_changed_files(filepaths)

    if not changed_paths:
        print("✅ All files unchanged — nothing to ingest.")
        return
    
    # Filter to only changed files
    files_to_process = [
        (name, text, fp)
        for name, text, fp in all_files
        if fp in changed_paths
    ]
    print(f"🔄 Processing {len(files_to_process)} changed file(s)...")


    all_docs = []
    all_embeddings = []
    all_ids = []
    all_metadata = []

    for filename, text, filepath in files_to_process:
        if not text.strip():
            print(f"Skipping empty files:{filename}")
            continue
        
        # Delete old chunks before re-ingesting
        delete_by_source(filename)
        
        # Use file type aware chunking
        suffix = Path(filepath).suffix.lower()
        if suffix == ".csv":
            chunks = chunk_csv(text)
        elif suffix == ".pdf":
            chunks = chunk_pdf(text)
        else:
            # TXT and DOCS both use standard recursive splitter
            chunks = chunk_text(text)

        embeddings = embed(chunks)


        # Safety check
        if len(chunks) != len(embeddings):
            raise ValueError(
                f"Embedding mismatch in {filename}:"
                f"{len(chunks)} chunks vs {len(embeddings)} embeddings"
            )

        for i, chunk in enumerate(chunks):
            doc_id = f"{filename}_{i}"

            all_docs.append(chunk)
            all_embeddings.append(embeddings[i])
            all_ids.append(doc_id)
            all_metadata.append({
                "source": filename,
                "chunk_id": i,
                "file_type": Path(filepath).suffix.lower().strip("."),
                "char_count": len(chunk)
            })
    # Store in chromaDB    
    collection.upsert(
        documents = all_docs,
        embeddings = all_embeddings,
        ids = all_ids,
        metadatas = all_metadata
    )

    save_bm25(all_docs, all_metadata)

    # Save updated hashes only after sucessful ingest
    save_hashes(updated_hashes)

    print(f"✅ Ingested {len(all_docs)} chunks from {len(files_to_process)} files")

# Run

if __name__ == "__main__":
    ingest()